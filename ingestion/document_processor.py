"""
Document processing module for PDFs and DOCs
"""
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Dict, List, Optional
from loguru import logger
import pytesseract
from PIL import Image
import io

from error_handler import (
    ErrorHandler, ErrorCategory, ErrorSeverity,
    handle_file_processing_error
)


class DocumentProcessor:
    """Handles extraction of text from PDF and DOC files"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
        self.error_handler = ErrorHandler()
    
    def process_file(self, file_path: Path) -> Dict:
        """
        Process a document file and extract text content with error handling
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing extracted text and metadata
        """
        def _process_with_validation():
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Validate file size
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            max_size_mb = 100  # From PROCESSING_CONFIG
            if file_size_mb > max_size_mb:
                raise ValueError(f"File too large: {file_size_mb:.1f}MB (max: {max_size_mb}MB)")
            
            suffix = file_path.suffix.lower()
            
            if suffix == '.pdf':
                return self._process_pdf(file_path)
            elif suffix in ['.docx', '.doc']:
                return self._process_docx(file_path)
            elif suffix == '.txt':
                return self._process_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {suffix}")
        
        return handle_file_processing_error(
            _process_with_validation,
            str(file_path),
            file_path.suffix.lower()
        )
    
    def _process_pdf(self, file_path: Path) -> Dict:
        """Extract text from PDF using PyMuPDF with error recovery"""
        doc = None
        try:
            doc = fitz.open(file_path)
            text_content = []
            total_pages = len(doc)
            
            for page_num in range(total_pages):
                try:
                    page = doc.load_page(page_num)
                    
                    # Extract text
                    text = page.get_text()
                    if text.strip():
                        text_content.append({
                            'page': page_num + 1,
                            'text': text.strip()
                        })
                    
                    # Extract images for OCR if text is sparse
                    if len(text.strip()) < 50:  # Likely image-heavy page
                        try:
                            pix = page.get_pixmap()
                            img_data = pix.tobytes("png")
                            img = Image.open(io.BytesIO(img_data))
                            
                            # Perform OCR with error handling
                            ocr_text = pytesseract.image_to_string(img)
                            if ocr_text.strip():
                                text_content.append({
                                    'page': page_num + 1,
                                    'text': ocr_text.strip(),
                                    'source': 'ocr'
                                })
                        except Exception as ocr_error:
                            logger.warning(f"OCR failed for page {page_num + 1}: {ocr_error}")
                            # Continue processing other pages
                            continue
                
                except Exception as page_error:
                    logger.warning(f"Failed to process page {page_num + 1}: {page_error}")
                    # Continue with next page
                    continue
            
            return {
                'file_path': str(file_path),
                'file_type': 'pdf',
                'content': text_content,
                'total_pages': total_pages,
                'processed_pages': len(text_content),
                'metadata': {
                    'title': doc.metadata.get('title', '') if doc.metadata else '',
                    'author': doc.metadata.get('author', '') if doc.metadata else '',
                    'subject': doc.metadata.get('subject', '') if doc.metadata else ''
                }
            }
            
        except Exception as e:
            error_report = self.error_handler.handle_error(
                e, ErrorCategory.FILE_PROCESSING,
                context={'file_path': str(file_path), 'file_type': 'pdf'},
                severity=ErrorSeverity.MEDIUM
            )
            
            # Try alternative PDF processing if available
            if error_report.recovery_action and 'alternative' in error_report.recovery_action:
                logger.info("Attempting alternative PDF processing method")
                try:
                    return self._process_pdf_alternative(file_path)
                except Exception as alt_error:
                    logger.error(f"Alternative PDF processing also failed: {alt_error}")
            
            raise
        finally:
            if doc:
                doc.close()
    
    def _process_docx(self, file_path: Path) -> Dict:
        """Extract text from DOCX files"""
        try:
            doc = Document(str(file_path))  # Convert Path to string
            text_content = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append({
                        'paragraph': i + 1,
                        'text': paragraph.text.strip()
                    })
            
            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_content.append({
                        'table': table_idx + 1,
                        'text': '\n'.join(table_text)
                    })
            
            return {
                'file_path': str(file_path),
                'file_type': 'docx',
                'content': text_content,
                'metadata': {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or ''
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _process_txt(self, file_path: Path) -> Dict:
        """Extract text from plain text files"""
        try:
            # Try different encodings to handle various text files
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            content = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        encoding_used = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError(f"Could not decode text file {file_path} with any supported encoding")
            
            # Split content into logical sections (paragraphs separated by double newlines)
            paragraphs = content.split('\n\n')
            text_content = []
            
            for i, paragraph in enumerate(paragraphs):
                cleaned_paragraph = paragraph.strip()
                if cleaned_paragraph:
                    text_content.append({
                        'paragraph': i + 1,
                        'text': cleaned_paragraph
                    })
            
            # If no double newlines, split by single newlines and group
            if len(text_content) <= 1 and content.strip():
                lines = content.split('\n')
                grouped_lines = []
                current_group = []
                
                for line in lines:
                    stripped_line = line.strip()
                    if stripped_line:
                        current_group.append(stripped_line)
                    else:
                        if current_group:
                            grouped_lines.append(' '.join(current_group))
                            current_group = []
                
                if current_group:
                    grouped_lines.append(' '.join(current_group))
                
                text_content = [
                    {'paragraph': i + 1, 'text': text}
                    for i, text in enumerate(grouped_lines)
                ]
            
            return {
                'file_path': str(file_path),
                'file_type': 'txt',
                'content': text_content,
                'metadata': {
                    'encoding': encoding_used,
                    'file_size': file_path.stat().st_size,
                    'line_count': len(content.split('\n'))
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise
    
    def _process_pdf_alternative(self, file_path: Path) -> Dict:
        """Alternative PDF processing using different library (placeholder)"""
        # This would use an alternative library like pdfplumber or PyPDF2
        # For now, return minimal structure to prevent complete failure
        logger.warning(f"Using fallback PDF processing for {file_path}")
        
        return {
            'file_path': str(file_path),
            'file_type': 'pdf',
            'content': [{'page': 1, 'text': 'Content extraction failed - file may be corrupted'}],
            'total_pages': 1,
            'processed_pages': 0,
            'metadata': {'processing_method': 'fallback'},
            'processing_error': True
        }
    
    def batch_process(self, file_paths: List[Path]) -> List[Dict]:
        """Process multiple documents with comprehensive error handling"""
        results = []
        failed_files = []
        
        for file_path in file_paths:
            try:
                result = self.process_file(file_path)
                if result:  # handle_file_processing_error might return None
                    results.append(result)
                    logger.info(f"Successfully processed: {file_path}")
                else:
                    failed_files.append(str(file_path))
                    logger.warning(f"Processing returned None for: {file_path}")
            except Exception as e:
                failed_files.append(str(file_path))
                error_report = self.error_handler.handle_error(
                    e, ErrorCategory.FILE_PROCESSING,
                    context={'file_path': str(file_path), 'batch_processing': True},
                    severity=ErrorSeverity.LOW  # Lower severity for batch processing
                )
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        # Log batch processing summary
        total_files = len(file_paths)
        successful_files = len(results)
        logger.info(f"Batch processing completed: {successful_files}/{total_files} files processed successfully")
        
        if failed_files:
            logger.warning(f"Failed files: {failed_files}")
        
        return results